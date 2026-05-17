"""DOCX export endpoints.

GET /api/export/cover-letter/{id}    — download cover letter as .docx
GET /api/export/interview-prep/{id}  — download interview prep sheet as .docx
"""

import io
import logging
from datetime import date
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.models.db import Application
from app.models.schemas import GapAnalysisOutput, InterviewPrepOutput

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/export", tags=["export"])

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docx_response(doc: Any, filename: str) -> StreamingResponse:
    """Serialise a python-docx Document to a StreamingResponse."""
    buffer = io.BytesIO()
    doc.save(buffer)  # type: ignore[attr-defined]
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type=_DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_cover_letter_docx(
    company: str,
    role_title: str,
    cover_letter_text: str,
) -> Any:
    """Build a professionally formatted cover letter DOCX."""
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(80)
        section.right_margin = Pt(80)

    # Header: role + company
    heading = doc.add_heading(f"{role_title}", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    sub = doc.add_paragraph(f"{company}")
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Date
    date_para = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    date_para.runs[0].font.size = Pt(11)
    date_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # spacer

    # Letter body — split on blank lines to preserve paragraph structure
    for block in cover_letter_text.split("\n\n"):
        block = block.strip()
        if block:
            p = doc.add_paragraph(block)
            p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(12)

    return doc


def _build_interview_prep_docx(
    company: str,
    role_title: str,
    gap_analysis: GapAnalysisOutput,
    interview_prep: InterviewPrepOutput,
) -> Any:
    """Build a formatted interview prep sheet DOCX."""
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(80)
        section.right_margin = Pt(80)

    # Title
    heading = doc.add_heading(f"Interview Prep — {role_title}", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    company_para = doc.add_paragraph(company)
    company_para.runs[0].font.size = Pt(12)
    company_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # Gap analysis summary
    doc.add_heading("Match Summary", level=2)

    score_para = doc.add_paragraph()
    score_run = score_para.add_run(f"Match Score: {gap_analysis.match_score}/100")
    score_run.bold = True
    score_run.font.size = Pt(11)

    reasoning_para = doc.add_paragraph(gap_analysis.match_reasoning)
    reasoning_para.runs[0].font.size = Pt(11)
    reasoning_para.paragraph_format.space_after = Pt(8)

    if gap_analysis.strengths:
        doc.add_paragraph("Key Strengths:").runs[0].bold = True
        for s in gap_analysis.strengths:
            p = doc.add_paragraph(s, style="List Bullet")
            p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Questions
    doc.add_heading("Interview Questions", level=2)

    category_labels = {
        "behavioral": "Behavioral",
        "technical": "Technical",
        "situational": "Situational",
        "culture_fit": "Culture Fit",
    }

    for i, q in enumerate(interview_prep.questions, start=1):
        label = category_labels.get(q.category, q.category.title())

        # Question line — bold
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i} [{label}]  {q.question}")
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_para.paragraph_format.space_before = Pt(10)

        # Framework line
        f_para = doc.add_paragraph()
        label_run = f_para.add_run("Suggested Framework:  ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        body_run = f_para.add_run(q.suggested_framework)
        body_run.font.size = Pt(10)
        f_para.paragraph_format.space_after = Pt(4)

    return doc


@router.get("/cover-letter/{app_id}")
async def export_cover_letter(
    app_id: int,
    db: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """Generate and return a DOCX file containing the cover letter for an application.

    Returns 404 if the application does not exist or has no cover letter text.
    """
    result = await db.execute(select(Application).where(Application.id == app_id))
    application = result.scalar_one_or_none()

    if application is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Application {app_id} not found.",
        )
    if not application.cover_letter_text:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Application {app_id} has no cover letter. Generate one first.",
        )

    doc = _build_cover_letter_docx(
        company=application.company,
        role_title=application.role_title,
        cover_letter_text=application.cover_letter_text,
    )

    slug = f"{application.company}-{application.role_title}".lower().replace(" ", "_")
    filename = f"cover_letter_{slug}.docx"

    logger.info("Exporting cover letter DOCX for application %d (%s)", app_id, application.company)
    return _docx_response(doc, filename)


@router.get("/interview-prep/{app_id}")
async def export_interview_prep(
    app_id: int,
    db: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """Return a DOCX interview prep sheet built from persisted interview_prep_json.

    Returns 404 if the application does not exist or has no saved interview prep.
    """
    result = await db.execute(select(Application).where(Application.id == app_id))
    application = result.scalar_one_or_none()

    if application is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Application {app_id} not found.",
        )
    if not application.interview_prep_json:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Application {app_id} has no interview prep. Generate one first.",
        )
    if not application.gap_analysis_json:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Application {app_id} has no gap analysis. Run /api/compare/resume first.",
        )

    prep = InterviewPrepOutput.model_validate_json(application.interview_prep_json)
    gap_analysis = GapAnalysisOutput.model_validate_json(application.gap_analysis_json)

    logger.info("Exporting interview prep DOCX for application %d (%s)", app_id, application.company)

    doc = _build_interview_prep_docx(
        company=application.company,
        role_title=application.role_title,
        gap_analysis=gap_analysis,
        interview_prep=prep,
    )

    slug = f"{application.company}-{application.role_title}".lower().replace(" ", "_")
    filename = f"interview_prep_{slug}.docx"

    return _docx_response(doc, filename)
